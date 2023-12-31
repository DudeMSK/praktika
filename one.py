import sys,os
import re
from collections import Counter
import win32com.client
import docx
from docx import Document
import textract
import shutil
from pathlib import Path
import glob
import docx2txt
from subprocess import Popen, PIPE
import pytesseract
import numpy as np
import sklearn
from sklearn.model_selection import train_test_split
from sklearn.naive_bayes import MultinomialNB
from sklearn.metrics import accuracy_score
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.svm import SVC

def find_uin_number(file_path):
    doc = Document(file_path)

    for para in doc.paragraphs:
        text = para.text
        match = re.search(r"УИН\s*№\s*(\d+)", text, re.UNICODE)
        if match:
            uin_number = match.group(1)
            return uin_number

    return None

#master_folder_path = r".\ТП_2021"
#source_folder_path = r".\ТП_2021"
#master_folder_path = r".\ТП_2022\ТП"
#source_folder_path = r".\ТП_2022\ТП"
master_folder_path = r".\ТП_2023\ТП"
source_folder_path = r".\ТП_2023\ТП"

#dest_notTXT = r".\ml\2021\notTXT"
#dest_TXT = r".\ml\2021\TXT"
#dest_notTXT = r".\ml\2022\notTXT"
#dest_TXT = r".\ml\2022\TXT"
dest_notTXT = r".\ml\2023\notTXT"
dest_TXT = r".\ml\2023\TXT"

pol_folder_path = os.path.join(dest_TXT, "pol")
otr_folder_path = os.path.join(dest_TXT, "otr")

for filename in os.listdir(source_folder_path):
    file_path = os.path.join(source_folder_path, filename)
    if file_path.endswith('.docx') and not filename.startswith('~$'):
        uin_number = find_uin_number(file_path)
        if uin_number:
            digits_only = re.sub(r"\D", "", uin_number)
            print(f"УИН № {digits_only} - {filename}")
            match = re.search(r'(\S+).*?(пол|отр)', filename)
            if match:
                numbers = match.group(1)
                word = match.group(2)
                print("Файл:", file_path)
                print("Категория заключения:", word)
                print("Номер реестра:", numbers)
                print()
                found_folders = []
                for root, dirs, files in os.walk(master_folder_path):
                    for dir_name in dirs:
                        if uin_number in dir_name:
                            folder_path = os.path.join(root, dir_name)
                            found_folders.append((uin_number, folder_path))
                if found_folders:
                    for uin, folder_path in found_folders:
                        print(f"Найдена папка с УИН № {uin} по пути: {folder_path}")
                        for root, dirs, files in os.walk(folder_path):
                            for file_name in files:
                                file_path = os.path.join(root, file_name)
                                if "Материалы по обоснованию в текстовой форме" in file_name:
                                    print(f"Найден файл 'Материалы по обоснованию в текстовой форме' по пути: {file_path}")
                                    dest_folder = dest_notTXT if file_path.endswith('.pdf') else dest_TXT
                                    new_file_name = f'{word} {numbers}{os.path.splitext(file_name)[1]}'
                                    dest_path = os.path.join(dest_folder, new_file_name)
                                    try:
                                        shutil.copy(file_path, dest_path)
                                        print(f"Файл {file_name} успешно скопирован в папку {dest_folder}!")
                                        new_file_path = os.path.join(dest_folder, new_file_name)
                                        os.rename(dest_path, new_file_path)
                                        print(f"Файл {dest_path} успешно переименован в {new_file_path}")
                                        if 'пол' in new_file_name:
                                            pol_folder = os.path.join(dest_folder, 'pol')
                                            if not os.path.exists(pol_folder):
                                                os.makedirs(pol_folder)
                                            shutil.move(new_file_path, pol_folder)
                                            print(f"Файл {new_file_name} перемещен в папку 'pol'")
                                        elif 'отр' in new_file_name:
                                            otr_folder = os.path.join(dest_folder, 'otr')
                                            if not os.path.exists(otr_folder):
                                                os.makedirs(otr_folder)
                                            shutil.move(new_file_path, otr_folder)
                                            print(f"Файл {new_file_name} перемещен в папку 'otr'")
                                    except Exception as e:
                                        print(f"Ошибка при копировании файла {file_name} в папку {dest_folder}: {str(e)}")
                else:
                    print(f"Папки с УИН {uin_number} не найдены в мастер-папке.")
                    print()
        else:
            print(f"Значение УИН не найдено в файле {filename}")
            print()
    else:
        
        print(f"Неверный формат файла: {filename}")

def convert_doc_to_docx(doc_file_path, docx_file_path):
    # Создаем объект приложения Word
    word_app = win32com.client.Dispatch("Word.Application")
    # Открываем документ .doc
    doc = word_app.Documents.Open(doc_file_path)
    # Сохраняем его в формате .docx
    doc.SaveAs(docx_file_path, 16)  # 16 означает формат .docx
    # Закрываем документ и выходим из приложения Word
    doc.Close()
    word_app.Quit()

# Конвертировать файлы .docx и .doc в папке "pol" в .txt
for new_file_name in os.listdir(pol_folder_path):
    file_path = os.path.join(pol_folder_path, new_file_name)
    print(f"{new_file_name} находится в {pol_folder_path}!")
    try:
        if file_path.endswith('.doc'):
            # Конвертировать .doc в .docx
            docx_filename = os.path.splitext(new_file_name)[0] + '.docx'
            docx_file_path = os.path.join(pol_folder_path, docx_filename)
            convert_doc_to_docx(file_path, docx_file_path)
            print(f"Файл {new_file_name} успешно сконвертирован в {docx_filename}!")
    except Exception as e:
        print(f"Ошибка при конвертации файла {new_file_name} в папку {docx_filename}: {str(e)}")

for new_file_name in os.listdir(pol_folder_path):
    file_path = os.path.join(pol_folder_path, new_file_name)
    print(f"{new_file_name} находится в {pol_folder_path}!")
    try:
        if file_path.endswith('.docx'):
            txt_filename = os.path.splitext(new_file_name)[0] + '.txt'
            txt_file_path = os.path.join(pol_folder_path, txt_filename)
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            with open(txt_file_path, "w", encoding="utf-8") as txt_file:
                txt_file.write(text)
            print(f"Файл {new_file_name} успешно сконвертирован в {txt_filename}!")
    except Exception as e:
        print(f"Ошибка при конвертации файла {new_file_name} в папку {pol_folder_path}: {str(e)}")

# Конвертировать файлы .docx и .doc в папке "otr" в .txt
for new_file_name in os.listdir(otr_folder_path):
    file_path = os.path.join(otr_folder_path, new_file_name)
    print(f"{new_file_name} находится в {otr_folder_path}!")
    try:
        if file_path.endswith('.doc'):
            # Конвертировать .doc в .docx
            docx_filename = os.path.splitext(new_file_name)[0] + '.docx'
            docx_file_path = os.path.join(otr_folder_path, docx_filename)
            convert_doc_to_docx(file_path, docx_file_path)
            print(f"Файл {new_file_name} успешно сконвертирован в {docx_filename}!")
    except Exception as e:
        print(f"Ошибка при конвертации файла {new_file_name} в папку {docx_filename}: {str(e)}")

for new_file_name in os.listdir(otr_folder_path):
    file_path = os.path.join(otr_folder_path, new_file_name)
    print(f"{new_file_name} находится в {otr_folder_path}!")
    try:
        if file_path.endswith('.docx'):
            txt_filename = os.path.splitext(new_file_name)[0] + '.txt'
            txt_file_path = os.path.join(otr_folder_path, txt_filename)
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            with open(txt_file_path, "w", encoding="utf-8") as txt_file:
                txt_file.write(text)
            print(f"Файл {new_file_name} успешно сконвертирован в {txt_filename}!")
    except Exception as e:
        print(f"Ошибка при конвертации файла {new_file_name} в папку {otr_folder_path}: {str(e)}")

# Код выполняет следующие действия:
# 1) find_uin_number(file_path): Это функция, которая принимает путь к файлу и ищет в нем УИН номер (уникальный идентификационный номер). Она открывает файл в формате .docx и проверяет каждый параграф на наличие соответствующего шаблона. Если номер найден, функция возвращает его, в противном случае возвращает None.
# 2) master_folder_path и source_folder_path: Это переменные, содержащие пути к мастер-папке и исходной папке соответственно. Вы можете изменить эти значения на свои пути.
# 3) dest_notTXT и dest_TXT: Это переменные, содержащие пути к папкам, куда будут копироваться файлы в зависимости от их формата. Файлы в формате .pdf будут скопированы в dest_notTXT, остальные файлы в dest_TXT. Вы также можете изменить эти значения по вашему усмотрению.
# 4) pol_folder_path и otr_folder_path: Это переменные, содержащие пути к папкам "pol" и "otr" внутри dest_TXT. В эти папки будут перемещены соответствующие файлы в зависимости от их категории.
# 5) Цикл for filename in os.listdir(source_folder_path): Этот цикл проходит через каждый файл в исходной папке. Для каждого файла проверяется, является ли он файлом формата .docx и не начинается ли имя файла с "~$". Если условие выполняется, выполняются дальнейшие действия.
# 6) uin_number = find_uin_number(file_path): Вызывается функция find_uin_number для поиска УИН номера в файле. Если номер найден, он сохраняется в переменной uin_number, иначе она принимает значение None.
# 7) digits_only = re.sub(r"\D", "", uin_number): Здесь происходит удаление всех символов, кроме цифр, из УИН номера, сохраненного в переменной uin_number. Результат сохраняется в переменной digits_only.
# 8) match = re.search(r'(\S+).*?(пол|отр)', filename): Эта строка ищет в имени файла соответствующий шаблон, состоящий из слова, за которым следует "пол" или "отр". Результат сохраняется в переменной match.
# 9) Цикл for root, dirs, files in os.walk(master_folder_path): Этот цикл рекурсивно обходит все папки и файлы в мастер-папке. Он ищет папку, содержащую УИН номер, и сохраняет путь к этой папке в переменной folder_path.
# 10) Цикл for file_name in files: Этот цикл проходит через все файлы в найденной папке.
# 11) if "Материалы по обоснованию в текстовой форме" in file_name: Это условие проверяет, содержит ли имя файла строку "Материалы по обоснованию в текстовой форме". Если условие выполняется, выполняются следующие действия.
# 12) dest_folder = dest_notTXT if file_path.endswith('.pdf') else dest_TXT: В зависимости от расширения файла (.pdf или другое), выбирается целевая папка для копирования.
# 13) new_file_name = f'{word} {numbers}{os.path.splitext(file_name)[1]}': Создается новое имя файла, состоящее из слова из шаблона, номера реестра и расширения файла.
# 14) shutil.copy(file_path, dest_path): Файл копируется из исходного пути file_path в целевую папку dest_path.
# 15) os.rename(dest_path, new_file_path): Файл переименовывается в новое имя new_file_path.
# 16) shutil.move(new_file_path, pol_folder): Файл перемещается в папку "pol".
# 17) shutil.move(new_file_path, otr_folder): Файл перемещается в папку "otr".
# 18) В остальной части кода выполняются аналогичные действия для остальных папок и файлов.

##############################################################################################
старое выполнение действий:
        
# 1. Определение функции find_uin_number, которая ищет номер УИН (уникального идентификационного номера) в переданном документе. Для этого функция открывает документ с помощью библиотеки docx, проходит по параграфам документа и ищет совпадения с помощью регулярного выражения.
# 2. Задание путей к папкам и файлам, которые будут использоваться в процессе. Например, master_folder_path - это путь к мастер-папке, source_folder_path - путь к папке, из которой будут браться файлы для обработки, и т.д.
# 3. Цикл for для перебора файлов в папке source_folder_path. Для каждого файла выполняются следующие шаги:
# 3.1 Проверка, что файл имеет расширение .docx и не является временным файлом (~$).
# 3.2 Поиск номера УИН в файле с помощью функции find_uin_number.
# 3.3 Если номер УИН найден, происходит извлечение информации о категории, номере реестра и т.д. из имени файла.
# 3.4 Поиск папок с соответствующими номерами УИН в мастер-папке.
# 3.5 Перебор файлов в найденных папках для выполнения дополнительных действий (например, копирования, переименования, перемещения файлов).
# 4. Определение функции convert_doc_to_docx, которая использует библиотеку win32com.client для конвертации документов .doc в .docx. Для этого функция открывает документ .doc с помощью приложения Word, сохраняет его в формате .docx и закрывает документ и приложение Word.
# 5. Цикл for для конвертации файлов .doc в .docx. Для каждого файла в папке pol_folder_path выполняются следующие шаги:
# 5.1 Проверка, что файл имеет расширение .doc.
# 5.2 Конвертация файла .doc в .docx с использованием функции convert_doc_to_docx.
# 6. Цикл for для конвертации файлов .docx в .txt. Для каждого файла в папке pol_folder_path выполняются следующие шаги:
# 6.1 Проверка, что файл имеет расширение .docx.
# 6.2 Открытие файла .docx с помощью библиотеки docx.
# 6.3 Извлечение текста из файла и сохранение его в формате .txt.

## В результате выполнения кода происходит обработка файлов в указанных папках.
## Он ищет номер УИН в документах, находит соответствующие папки в мастер-папке, копирует, переименовывает и перемещает файлы, а также конвертирует документы .doc в .docx и .docx в .txt.
